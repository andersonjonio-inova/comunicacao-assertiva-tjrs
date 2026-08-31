from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE = Path(__file__).resolve().parent
OUT = BASE / "entregas" / "Estudo_de_Caso_O_Relatorio_Que_Nao_Ficara_Pronto_Versao_Final.docx"
LOGO = BASE / "assets" / "logo-cjud.jpg"

BLUE = "17488F"
NAVY = "0E2F66"
GREEN = "76B82A"
ORANGE = "F39A24"
YELLOW = "FFD21D"
INK = "15243A"
MUTED = "617087"
WASH = "F4F8FC"
LINE = "DCE5EF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char_1, instr_text, fld_char_2))


def add_color_signature(doc):
    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    for cell, color in zip(table.rows[0].cells, (BLUE, GREEN, ORANGE, YELLOW)):
        shade(cell, color)
        set_cell_margins(cell, top=45, start=0, bottom=45, end=0)


def add_cover(doc):
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_p.add_run().add_picture(str(LOGO), width=Cm(2.45))

    institution = doc.add_paragraph()
    institution.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = institution.add_run("TJRS · FORMAÇÃO DE GESTORES")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    banner = doc.add_table(rows=1, cols=1)
    cell = banner.cell(0, 0)
    shade(cell, NAVY)
    set_cell_margins(cell, top=360, start=420, bottom=360, end=420)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("COMUNICAÇÃO ASSERTIVA")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(255, 255, 255)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Inteligência Emocional e Comunicação Não Violenta")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string("DDEAFF")

    add_color_signature(doc)
    doc.add_paragraph()

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("O relatório que não ficará pronto")
    run.bold = True
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Estudo de caso em duas partes | Inteligência emocional aplicada à gestão")
    run.italic = True
    run.font.color.rgb = RGBColor.from_string(MUTED)

    formadores = doc.add_table(rows=2, cols=1)
    formadores.autofit = True
    shade(formadores.cell(0, 0), WASH)
    shade(formadores.cell(1, 0), "FFFFFF")
    for cell in formadores.column_cells(0):
        set_cell_margins(cell, top=150, start=240, bottom=150, end=240)
    p = formadores.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FORMADORES")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(GREEN)
    p = formadores.cell(1, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Ânderson Porto  •  Fernando de Assis Alves")
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = RGBColor.from_string(NAVY)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = intro.add_run("Caso fictício para uso formativo. Leia apenas a parte indicada pelo facilitador.")
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_page_break()


def configure_header_footer(section):
    section.different_first_page_header_footer = True
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(16.8))
    table.columns[0].width = Cm(1.25)
    table.columns[1].width = Cm(15.55)
    logo_cell, text_cell = table.rows[0].cells
    logo_cell.paragraphs[0].add_run().add_picture(str(LOGO), width=Cm(0.72))
    p = text_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("COMUNICAÇÃO ASSERTIVA · TJRS")
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p = text_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Inteligência Emocional e Comunicação Não Violenta")
    r.font.size = Pt(7.5)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Ânderson Porto e Fernando de Assis Alves   •   ")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_number(footer)


def add_case_box(doc, title, text):
    table = doc.add_table(rows=2, cols=1)
    table.autofit = True
    shade(table.cell(0, 0), NAVY)
    shade(table.cell(1, 0), "FFF8ED")
    for cell in table.column_cells(0):
        set_cell_margins(cell)
    p = table.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    r.font.size = Pt(12)
    p = table.cell(1, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10.5)
    doc.add_paragraph()


def add_question_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Cm(6.2)
    table.columns[1].width = Cm(10.2)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for cell, text in zip(hdr.cells, ("Pergunta de análise", "Registro do grupo")):
        shade(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
    for index, prompt in enumerate(rows):
        cells = table.add_row().cells
        cells[0].text = prompt
        cells[1].text = "\n\n\n"
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if index % 2 == 0:
                shade(cell, WASH)
        cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(INK)
        cells[0].paragraphs[0].runs[0].bold = True
    doc.add_paragraph()


def add_body_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.12
    return p


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    configure_header_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    for style_name, size, color in (
        ("Title", 24, NAVY),
        ("Heading 1", 17, BLUE),
        ("Heading 2", 13, ORANGE),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")

    add_cover(doc)

    doc.add_heading("Orientação para o trabalho em grupo", level=1)
    add_body_paragraph(doc, "O caso foi construído com informações deliberadamente incompletas. O objetivo não é decidir rapidamente quem está certo, diagnosticar pessoas ou oferecer uma solução imediata. A tarefa é observar como fatos, interpretações, estados corporais, emoções, impulsos e escolhas comunicacionais se relacionam em uma situação de gestão.")
    add_body_paragraph(doc, "Na primeira etapa, trabalhe somente com a Parte 1. Após a discussão e a orientação do facilitador, avance para a Parte 2. Ao formular hipóteses sobre a experiência interna dos personagens, use expressões como “pode ter sentido”, “é possível que tenha interpretado” ou “seria necessário verificar”.")

    doc.add_heading("Parte 1 | A conversa no final do expediente", level=1)
    doc.add_heading("Cena 1: prazo, pressão e uma resposta pública", level=2)
    add_body_paragraph(doc, "É quinta-feira, 16h40. Em uma unidade do Tribunal, a equipe está encerrando uma semana marcada por demandas simultâneas e mudanças de prioridade. Na manhã seguinte, a chefia deverá encaminhar à direção um relatório consolidado, que servirá de apoio para uma reunião já agendada. O gestor considera a entrega importante e acompanha o prazo com preocupação.")
    add_body_paragraph(doc, "O servidor responsável pela consolidação trabalhou no relatório ao longo dos últimos dias. Naquela tarde, porém, comunica ao gestor que o documento não ficará pronto no prazo previsto. Explica que recebeu outra demanda urgente da direção e precisou interromper a tarefa. Acrescenta que parte dos dados esperados de outras áreas chegou incompleta e que ainda seria necessário conferir algumas informações antes de fechar a versão final.")
    add_body_paragraph(doc, "Enquanto escuta, o gestor percebe o horário e olha para o relógio. Sua respiração fica mais curta, o peito parece mais apertado e a musculatura dos ombros e da mandíbula se contrai. A energia corporal aumenta. Sua atenção se estreita em torno de três elementos: o horário, a frase “não ficará pronto” e a possibilidade de comparecer à reunião da manhã seguinte sem o documento. Quase ao mesmo tempo, surgem pensamentos como: “De novo vou ter de responder por uma entrega que não aconteceu”, “a direção vai entender que eu não controlo a equipe” e “se eu aceitar essa explicação, o prazo deixará de ser levado a sério”.")
    add_body_paragraph(doc, "A experiência parece desagradável e de alta ativação. Preocupação, irritação e medo de exposição podem estar misturados. O impulso mais imediato é interromper a explicação, reafirmar autoridade e pressionar por uma solução. Ainda não sabemos se o gestor reconhece esses sinais como informação sobre seu próprio estado, nem se distingue o risco objetivo da interpretação que está construindo naquele instante.")
    add_body_paragraph(doc, "Antes que o servidor conclua a explicação, o gestor o interrompe diante de dois colegas que permanecem na sala e afirma, em tom firme:")
    add_case_box(doc, "FALA DO GESTOR", "“Isso é falta de comprometimento. Toda vez é a mesma coisa. Não quero justificativas. Dê um jeito e entregue amanhã cedo.”")
    add_body_paragraph(doc, "A sala fica em silêncio. O servidor não responde imediatamente. Os dois colegas desviam o olhar. Nesse ponto, a cena é interrompida para análise.")

    doc.add_heading("Mapa do Caso | Primeira análise", level=2)
    add_question_table(doc, [
        "1. O que pode ser observado? Quais fatos uma câmera, uma gravação ou um registro documental poderia confirmar?",
        "2. O que foi interpretado? Quais avaliações, generalizações, suposições ou conclusões aparecem como se fossem fatos?",
        "3. Quais necessidades parecem estar mobilizadas no gestor e no servidor? Que emoções podem emergir quando essas necessidades são atendidas, frustradas ou percebidas como ameaçadas?",
    ])

    doc.add_page_break()
    doc.add_heading("Parte 2 | O impacto da fala", level=1)
    doc.add_heading("Cena 2: da exposição pública ao transbordamento emocional", level=2)
    add_body_paragraph(doc, "Nos segundos seguintes à fala do gestor, o servidor sente o rosto aquecer e o coração acelerar. A garganta parece fechar, a respiração fica presa por alguns instantes e surge tensão no peito e no abdômen. Ele aperta os lábios e mantém as mãos imóveis sobre a mesa para não responder no impulso. A experiência é intensamente desagradável e sua ativação sobe rapidamente.")
    add_body_paragraph(doc, "Ao notar que os colegas presenciaram a repreensão, sua atenção se desloca do problema técnico do relatório para a própria exposição. Ele percebe o silêncio da sala, os olhares desviados e o tom usado pelo gestor. Surgem pensamentos em sequência: “Ele nem me deixou terminar”, “está dizendo diante de todos que eu não sou comprometido”, “não importa o que eu faça” e “talvez meus colegas pensem que sou incompetente”. Vergonha, raiva, frustração, injustiça percebida e medo de desvalorização podem coexistir. O impulso oscila entre se defender, atacar, abandonar a conversa e permanecer calado para evitar uma escalada.")
    add_body_paragraph(doc, "O servidor responde apenas: “Entendi”. Abre novamente o arquivo, mas encontra dificuldade para organizar as ideias. Lê o mesmo parágrafo mais de uma vez, comete pequenos erros e alterna entre o documento e a lembrança da cena. Parte de sua energia mental passa a ser consumida por uma conversa imaginária na qual explica tudo o que não conseguiu dizer. A ativação corporal diminui um pouco, mas não desaparece. Cada vez que recorda a expressão “falta de comprometimento”, sente novo calor no rosto e volta a contrair a mandíbula.")
    add_body_paragraph(doc, "Ao sair do trabalho, ele leva consigo a preocupação com o prazo e a sensação de ter sido publicamente desconsiderado. No caminho para casa, repassa o episódio, imagina o que os colegas podem comentar e ensaia respostas que gostaria de ter dado. A ruminação mantém a experiência ativa mesmo sem a presença do gestor. Em alguns momentos predomina a raiva, com energia para confronto; em outros, a ativação cai e aparecem desânimo, cansaço e vontade de se afastar.")
    add_body_paragraph(doc, "Em casa, sua esposa pergunta como foi o dia. Ele responde “normal” e evita contar o que aconteceu. Pouco depois, ela pergunta: “Você conseguiu resolver o problema do carro?”. A pergunta é cotidiana, mas ele a recebe quando ainda está emocionalmente mobilizado. Interpreta-a como mais uma cobrança, eleva o tom de voz e responde: “Eu não quero ser cobrado”. A esposa se surpreende e se cala. A situação mudou, mas parte da ativação, da atenção defensiva e dos significados construídos no trabalho alcançou outra relação.")
    add_body_paragraph(doc, "Na manhã seguinte, o relatório e as responsabilidades funcionais continuam precisando ser tratados. Agora, porém, existe também um segundo problema: a qualidade da relação entre gestor e servidor, a segurança para comunicar riscos futuros e os efeitos do episódio sobre concentração, confiança e cooperação na equipe.")

    doc.add_heading("Mapa do Caso | Segunda análise", level=2)
    add_question_table(doc, [
        "1. Que mudanças corporais o servidor percebeu? Como a interocepção poderia ajudá-lo a reconhecer o próprio estado antes de agir?",
        "2. Como valência e ativação variaram entre a fala do gestor, o trabalho após a cena, o trajeto e a chegada em casa?",
        "3. Quais emoções, pensamentos e impulsos podem ser distinguidos? Quais deles são fatos e quais são interpretações ou hipóteses?",
        "4. Onde aparecem ruminação, transferência de ativação e transbordamento emocional?",
        "5. Que respostas poderiam preservar, ao mesmo tempo, a responsabilidade pela entrega, a dignidade das pessoas e a qualidade da relação de trabalho?",
    ])

    doc.add_heading("Síntese conceitual para consulta", level=1)
    concepts = [
        ("Interocepção", "Percepção dos sinais internos do organismo, como batimentos, respiração, temperatura, tensão, desconforto visceral e alterações de energia. Esses sinais informam sobre o estado corporal, mas não determinam uma emoção ou conduta isoladamente."),
        ("Valência", "Qualidade agradável, desagradável ou neutra da experiência. Situações semelhantes podem ter valências diferentes conforme contexto, história e interpretação."),
        ("Ativação", "Nível de mobilização do organismo, de baixo a alto. Alta ativação pode estreitar a atenção e favorecer respostas rápidas; baixa ativação pode aparecer como desânimo, retraimento ou redução de energia."),
        ("Inteligência emocional na gestão", "Capacidade de perceber, utilizar, compreender e administrar informações emocionais, articulando autoconsciência, autorregulação, empatia e habilidades relacionais sem abandonar critérios, limites e responsabilidade profissional."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for cell, text in zip(hdr.cells, ("Conceito", "Aplicação no caso")):
        shade(cell, NAVY)
        r = cell.paragraphs[0].add_run(text)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
    for name, explanation in concepts:
        cells = table.add_row().cells
        cells[0].text = name
        cells[1].text = explanation
        cells[0].paragraphs[0].runs[0].bold = True
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    doc.core_properties.title = "O relatório que não ficará pronto"
    doc.core_properties.subject = "Estudo de caso em duas partes sobre inteligência emocional aplicada à gestão"
    doc.core_properties.author = "Ânderson Porto e Fernando de Assis Alves"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
